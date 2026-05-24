from pathlib import Path

import fitz
import pandas as pd
from docx import Document


def extract_text(file_path):
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in [".xlsx", ".xls"]:
        return extract_excel(path)
    if suffix == ".csv":
        return extract_csv(path)
    if suffix in [".txt", ".md"]:
        return path.read_text(encoding="utf-8", errors="ignore")

    raise ValueError(f"Unsupported file format: {suffix}")


def extract_pdf(path):
    parts = []
    with fitz.open(path) as doc:
        for page_number, page in enumerate(doc, start=1):
            text = page.get_text()
            if text.strip():
                parts.append(f"\n[Page {page_number}]\n{text}")
    return "\n".join(parts)


def extract_docx(path):
    document = Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    table_parts = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_parts.append(" | ".join(cells))

    return "\n".join(paragraphs + table_parts)


def extract_excel(path):
    sheets = pd.read_excel(path, sheet_name=None)
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"\n[Sheet: {sheet_name}]\n")
        parts.append(df.fillna("").to_csv(index=False))
    return "\n".join(parts)


def extract_csv(path):
    df = pd.read_csv(path)
    return df.fillna("").to_csv(index=False)